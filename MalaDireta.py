from docx import Document
from docx.shared import Pt
from datetime import datetime
import tempfile
import os
from docx2pdf import convert


# ----------------------------------------------
# Substituir TAGS
# ----------------------------------------------
def substituir_tags(paragrafo, tags: dict):
    for tag, valor in tags.items():
        if tag in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(tag, str(valor))
    # aplica Arial 8
    for run in paragrafo.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)


# ----------------------------------------------
# GERAR PDF FINAL DA PROPOSTA
# ----------------------------------------------
def gerar_documento_proposta_pdf(supabase, id_proposta, caminho_template, nome_pdf):
    """
    Gera apenas PDF (nenhum DOCX é salvo).
    1. Carrega matriz.docx
    2. Preenche tags
    3. Preenche tabela 2 com itens
    4. Converte para PDF
    5. Retorna caminho final do PDF
    """

    # ------------------------------------------
    # Buscar dados no Supabase
    # ------------------------------------------
    proposta = (
        supabase.table("propostas")
        .select("*")
        .eq("id_proposta", id_proposta)
        .single()
        .execute()
    ).data

    cliente = (
        supabase.table("clientes")
        .select("*")
        .eq("id", proposta["id_cliente"])
        .single()
        .execute()
    ).data

    itens = (
        supabase.table("itens_proposta")
        .select("*")
        .eq("id_proposta", id_proposta)
        .execute()
    ).data

    # ------------------------------------------
    # Carregar template
    # ------------------------------------------
    doc = Document(caminho_template)

    # ------------------------------------------
    # Criar dicionário de TAGS
    # ------------------------------------------
    data_emissao_formatada = datetime.strptime(
        proposta["data_emissao"], "%Y-%m-%d"
    ).strftime("%d/%m/%Y")

    tags = {
        "{{NUM_PROPOSTA}}": proposta["num_proposta"],
        "{{DATA_EMISSAO}}": data_emissao_formatada,
        "{{VALIDADE}}": proposta["validade"],
        "{{COND_PAGAMENTO}}": proposta["cond_pagamento"],
        "{{REFERENCIA}}": proposta["referencia"],

        "{{ID}}": cliente["id"],
        "{{CLIENTE}}": cliente["empresa"],
        "{{CNPJ}}": cliente["cnpj"],
        "{{ENDERECO}}": cliente["endereco"],
        "{{CIDADE}}": cliente["cidade"],
        "{{UF}}": cliente["uf"],
        "{{CONTATO}}": cliente["contato"],
        "{{EMAIL}}": cliente["email"],
        "{{TELEFONE}}": cliente["telefone"],
    }

    # ------------------------------------------
    # Substituir tags nos parágrafos
    # ------------------------------------------
    for p in doc.paragraphs:
        substituir_tags(p, tags)

    # ------------------------------------------
    # Substituir tags nas tabelas
    # ------------------------------------------
    for tabela in doc.tables:
        for linha in tabela.rows:
            for cel in linha.cells:
                for p in cel.paragraphs:
                    substituir_tags(p, tags)

    # ==================================================
    # TABELA DE ITENS — segunda tabela do documento
    # ==================================================
    tabela_itens = doc.tables[1]  # <-- IMPORTANTE!

    # Garantir 8 colunas
    if len(tabela_itens.rows[0].cells) != 8:
        raise ValueError("A segunda tabela do template deve possuir 8 colunas.")

    # Remover linhas extras, mantendo somente o cabeçalho
    while len(tabela_itens.rows) > 1:
        tabela_itens._element.remove(tabela_itens.rows[-1]._element)

    total_proposta = 0

    # ------------------------------------------
    # Inserir itens
    # ------------------------------------------
    for idx, item in enumerate(itens, start=1):
        qtd = float(item["qtd"])
        preco = float(item["preco_unitario"])
        desconto = float(item["desconto"])
        total_item = (qtd * preco) - (preco * desconto / 100)
        total_proposta += total_item

        linha = tabela_itens.add_row().cells

        linha[0].text = str(idx)
        linha[1].text = item["codigo_servico"]
        linha[2].text = item["descricao_servico"]
        linha[3].text = item["prazo_ddl"]
        linha[4].text = str(int(qtd))
        linha[5].text = f"R$ {preco:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        linha[6].text = f"{desconto:.2f}%"
        linha[7].text = f"R$ {total_item:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        # Arial 8
        for cel in linha:
            for p in cel.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)

    # Linha em branco
    linha_vazia = tabela_itens.add_row().cells
    for cel in linha_vazia:
        r = cel.paragraphs[0].add_run(" ")
        r.font.name = "Arial"
        r.font.size = Pt(8)

    # Sub.Total
    linha_total = tabela_itens.add_row().cells
    linha_total[5].text = "Sub.Total:"
    linha_total[7].text = (
        f"R$ {total_proposta:,.2f}"
        .replace(",", "X").replace(".", ",").replace("X", ".")
    )

    for cel in linha_total:
        for p in cel.paragraphs:
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8)

    # ======================================================
    #  SALVAR EM TEMPORÁRIO COMO DOCX → CONVERTER PARA PDF
    # ======================================================
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    temp_pdf  = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name

    doc.save(temp_docx)
    convert(temp_docx, temp_pdf)

    # ======================================================
    #  RENOMEAR PARA O NOME ESCOLHIDO PELO USUÁRIO
    # ======================================================
    caminho_pdf_final = os.path.join(tempfile.gettempdir(), nome_pdf)
    os.replace(temp_pdf, caminho_pdf_final)

    return caminho_pdf_final
